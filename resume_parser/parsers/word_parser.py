"""Word document parser implementation."""

import os
import logging
from docx import Document

logger = logging.getLogger(__name__)


class WordParser:
    """Parser for Word documents using python-docx."""
    
    def parse(self, file_path: str) -> str:
        """Extract text from a Word document."""
        logger.info(f"Starting Word document parsing for: {file_path}")
        
        if not os.path.exists(file_path):
            logger.error(f"Word document not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract from paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
                    paragraph_count += 1
            
            logger.debug(f"Extracted text from {paragraph_count} paragraphs")
            
            # Extract from tables
            table_count = 0
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
                        table_count += 1
            
            if table_count > 0:
                logger.debug(f"Extracted text from {table_count} table rows")
            
            if not text_content:
                logger.error(f"No text content could be extracted from Word document: {file_path}")
                raise ValueError("No text content could be extracted from Word document")
            
            total_chars = len("\n".join(text_content))
            logger.info(f"Successfully extracted {total_chars} characters from Word document")
            
            return "\n".join(text_content)
            
        except Exception as e:
            if "No text content could be extracted" in str(e):
                raise  # Re-raise our custom error
            logger.error(f"Error parsing Word document {file_path}: {e}")
            raise